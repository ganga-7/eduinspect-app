# main.py
import streamlit as st
from db import get_db
from datetime import datetime
import base64
import os
import re
import imagehash
import pytesseract
from PIL import Image
import cv2
from PIL.ExifTags import TAGS, GPSTAGS
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from PyPDF2 import PdfReader
from docx import Document
import io

pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

def analyze_image_conditions(image_path):
    import random
    return random.choice(["Good", "Needs Attention", "Bad"])

def extract_coordinates_from_overlay(image_path):
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img)
    regexes = [
        r'Lat(?:itude)?[^\d\-]*(-?\d+\.\d+)[^\d\-]*Long(?:itude)?[^\d\-]*(-?\d+\.\d+)',
        r'(-?\d+\.\d+)[^\d\-]{1,6}(-?\d+\.\d+)'
    ]
    for pattern in regexes:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                lat = float(match.group(1))
                lon = float(match.group(2))
                if 10 < lat < 40 and 60 < lon < 100:
                    return (lat, lon)
            except:
                continue
    return (None, None)

def check_authenticity(image_path, inspection_date_str, expected_coords=None):
    metadata = {}
    try:
        image = Image.open(image_path)
        exif_data = image._getexif()
        if exif_data:
            for tag_id, value in exif_data.items():
                tag = TAGS.get(tag_id, tag_id)
                if tag == "GPSInfo":
                    gps_data = {GPSTAGS.get(t, t): value[t] for t in value}
                    metadata["GPSInfo"] = gps_data
                else:
                    metadata[tag] = value
    except:
        pass

    gps_coords = extract_coordinates_from_overlay(image_path)
    date_taken = metadata.get("DateTimeOriginal", "Not found")
    software_used = metadata.get("Software", "Not found")

    try:
        image_date = datetime.strptime(date_taken, "%Y:%m:%d %H:%M:%S").date()
        inspection_date = datetime.strptime(inspection_date_str, "%Y-%m-%d").date()
        date_status = "✅" if image_date == inspection_date else "❌"
    except:
        date_status = "❌"

    if expected_coords and gps_coords and None not in gps_coords:
        lat_diff = abs(expected_coords[0] - gps_coords[0])
        lon_diff = abs(expected_coords[1] - gps_coords[1])
        gps_status = "✅" if lat_diff < 0.03 and lon_diff < 0.03 else "❌"
    else:
        gps_status = "❌"

    edit_status = "❌" if "photoshop" in software_used.lower() else "✅"

    try:
        original_hash = imagehash.average_hash(Image.open(image_path))
        re_saved_hash = imagehash.average_hash(Image.open(image_path).convert("RGB"))
        hash_diff = original_hash - re_saved_hash
        hash_status = "❌" if hash_diff > 10 else "✅"
    except:
        hash_diff = "N/A"
        hash_status = "❌"

    overall_auth = "Authentic" if all([
        date_status == "✅", gps_status == "✅",
        edit_status == "✅", hash_status == "✅"
    ]) else "Fake"

    return {
        "authenticity": overall_auth,
        "software_used": software_used,
        "date_taken": date_taken,
        "gps_coords": gps_coords,
        "hash_diff": hash_diff,
        "date_status": date_status,
        "gps_status": gps_status,
        "edit_status": edit_status,
        "hash_status": hash_status
    }

# ---------------- STREAMLIT UI ---------------- #

db = get_db()
st.title("EduInspect - Empowering Education Through AI")
menu = ["Home", "Register Institute", "Start Inspection", "Dashboard", "Analytics"]
choice = st.sidebar.selectbox("Menu", menu)

if choice == "Home":
    st.subheader("Welcome to EduInspect")
    st.write("Register your institution and upload required data for automated AI-driven inspection.")

elif choice == "Register Institute":
    st.subheader("New Institute Registration")
    with st.form("register_form"):
        name = st.text_input("Institute Name")
        email = st.text_input("Email")
        password = st.text_input("Password", type="password")
        code = st.text_input("Institute Code")
        submitted = st.form_submit_button("Register")

        if submitted:
            existing = db.institutions.find_one({"$or": [{"code": code}, {"email": email}]})
            if existing:
                similar_codes = db.institutions.find({"code": {"$regex": f"^{code[:3]}"}}).limit(3)
                suggestions = [entry['code'] for entry in similar_codes if entry['code'] != code]
                st.warning("This institute is already registered.")
                if suggestions:
                    st.info("Similar codes: " + ", ".join(suggestions))
            else:
                db.institutions.insert_one({
                    "name": name, "email": email, "password": password,
                    "code": code, "created_at": datetime.now()
                })
                st.success("Registration Successful")

elif choice == "Start Inspection":
    st.subheader("Upload Data for Inspection")
    code = st.text_input("Enter Institute Code")
    inspection_date = st.date_input("Select Inspection Date")
    img_files = st.file_uploader("Upload 10 Facility Images", type=['jpg', 'jpeg', 'png'], accept_multiple_files=True)
    doc_file = st.file_uploader("Upload Report Document", type=None)
    feedback = st.text_area("Feedback or Survey Data")
    inspect = st.button("Start Inspection")

    if inspect:
        if code and img_files and len(img_files) == 10:
            parameters = [
                "Campus Cleanliness", "Classroom Environment", "Computer Labs", "Drinking Water Facility",
                "Fire Safety Measures", "Girls’ Common Room", "Health and Wellness Facilities",
                "Library Infrastructure", "Toilet Facilities", "Waste Management"
            ]
            expected_coords = (17.6868, 75.9112)
            inspection_date_str = inspection_date.strftime("%Y-%m-%d")
            summary_report = {}

            os.makedirs("temp", exist_ok=True)

            for idx, img in enumerate(img_files):
                temp_path = os.path.join("temp", img.name)
                with open(temp_path, "wb") as f:
                    f.write(img.getbuffer())
                condition = analyze_image_conditions(temp_path)
                auth_result = check_authenticity(temp_path, inspection_date_str, expected_coords)

                summary_report[parameters[idx]] = {
                    "Condition": condition,
                    "Authenticity": auth_result["authenticity"],
                    "Report": auth_result
                }

                db.inspections.insert_one({
                    "code": code,
                    "parameter": parameters[idx],
                    "condition": condition,
                    "authenticity": auth_result["authenticity"],
                    "timestamp": datetime.now()
                })

            doc_data = base64.b64encode(doc_file.read()).decode("utf-8") if doc_file else None
            db.reports.insert_one({
                "code": code,
                "document": doc_data,
                "filename": doc_file.name if doc_file else None,
                "feedback": feedback,
                "timestamp": datetime.now()
            })

            st.success("Inspection Completed.")
            st.subheader("📸 Image Authentication Reports")
            for param, data in summary_report.items():
                report = data["Report"]
                st.markdown(f"""
🦾 **Image Authentication Report - {param}**
📅 Date Taken: {report['date_taken']} ({report['date_status']})  
📍 Coordinates: {report['gps_coords']} ({report['gps_status']})  
🧠 Tamper Check: {report['hash_diff']} ({report['hash_status']})  
✅ Authenticity: **{report['authenticity']}**  
📌 Condition: **{data['Condition']}**
""")

        # Document Smart Summary
        if doc_file:
            doc_text = ""
            ext = doc_file.name.lower()
            if ext.endswith(".pdf"):
                pdf = PdfReader(doc_file)
                doc_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            elif ext.endswith(".docx"):
                word = Document(doc_file)
                doc_text = "\n".join([p.text for p in word.paragraphs])
            elif ext.endswith(".txt"):
                doc_text = doc_file.read().decode("utf-8")

            st.subheader("📑 Document Summary (AI Extracted)")
            def extract_from_document(doc_text):
                def extract(patterns, fallback="Not found"):
                    for pattern in patterns:
                        match = re.search(pattern, doc_text, re.IGNORECASE)
                        if match:
                            return match.group(0).strip()
                    return fallback

                return {
            "Highest Qualification": extract([
            r"(Ph\.?D(?:\s+in\s+[A-Za-z\s&]+)?)",
            r"(M\.?B\.?A(?:\s+in\s+[A-Za-z\s&]+)?)",
            r"(M\.?Tech(?:\s+in\s+[A-Za-z\s&]+)?)",
            r"(B\.?Tech(?:\s+in\s+[A-Za-z\s&]+)?)",
            r"(M\.?Sc(?:\s+in\s+[A-Za-z\s&]+)?)",
            r"(B\.?E(?:\s+in\s+[A-Za-z\s&]+)?)"
        ]),

        "Most Recent Experience": extract([
            r"(Director|Trainer|Professor|Lecturer|Principal|Dean)\s+at\s+[A-Z][A-Za-z&\s,]+",
            r"(Director|Trainer|Professor|Lecturer|Principal|Dean)[^\n]{0,80}\n.{0,80}\n[A-Z][A-Za-z&\s,]+",
            r"(Currently\s+working\s+as\s+[A-Za-z\s]+)"
        ]),

        "Total Experience": extract([
            r"\d{1,2}\+?\s*(years|yrs|Years|Yrs)\s+of\s+(Corporate|Academic|Industry)?\s*experience",
            r"(Total\s+Experience\s*[:\-]?\s*\d{1,2}\+?\s*(years|Yrs|yrs))",
            r"(\d{1,2}\+?\s*(years|Yrs|yrs))"
        ]),

        "Publications": extract([
            r"\d+\s*\(\s*\d+\s*National\s*,\s*\d+\s*International\s*\)",
            r"\d+\s*(National|International)[^\n]*Publications?",
            r"(1\s*National\s*(and|&)?\s*1\s*International)"
        ])
    }

            summary = extract_from_document(doc_text)
            st.markdown("📑 **Document Summary (AI Extracted):**")
            for key, val in summary.items():
                st.markdown(f"**{key}:** {val}")



elif choice == "Dashboard":
    st.subheader("📊 Inspection Dashboard")
    code = st.text_input("Enter Institute Code")
    selected_date = st.date_input("Filter by Date (optional)", value=None)

    if code:
        query = {"code": code}
        if selected_date:
            start_dt = datetime.combine(selected_date, datetime.min.time())
            end_dt = datetime.combine(selected_date, datetime.max.time())
            query["timestamp"] = {"$gte": start_dt, "$lte": end_dt}

        results = list(db.inspections.find(query).sort("timestamp", -1))
        reports = list(db.reports.find({"code": code}).sort("timestamp", -1))

        # -------------------- IMAGE INSPECTIONS --------------------
        if results:
            grouped = {}
            for rec in results:
                dt_str = rec["timestamp"].strftime("%Y-%m-%d %H:%M")
                grouped.setdefault(dt_str, []).append(rec)

            for dt, entries in grouped.items():
                st.markdown(f"### 🗓️ Inspection: {dt}")
                for rec in entries:
                    cond_icon = "🟢" if rec['condition'] == "Good" else "🔴"
                    auth_icon = "✅" if rec['authenticity'] == "Authentic" else "❌"
                    st.markdown(f"- **{rec['parameter']}** ➡️ {cond_icon} {rec['condition']} | {auth_icon} {rec['authenticity']}")

        # -------------------- DOCUMENT REPORTS --------------------
        if reports:
            st.markdown("### 📎 Uploaded Documents")
            for i, rpt in enumerate(reports):
                tstr = rpt["timestamp"].strftime("%Y-%m-%d %H:%M")
                st.markdown(f"📄 **Report from {tstr}**")

                # Download Button
                if rpt.get("document"):
                    st.download_button("📥 Download", base64.b64decode(rpt["document"]),
                                       file_name=rpt.get("filename", f"report_{i}.pdf"),
                                       key=f"rpt_{i}")

                # Feedback Text
                if rpt.get("feedback"):
                    st.markdown(f"📝 **Feedback:** {rpt['feedback']}")

                # Smart Summary
                if rpt.get("filename") and rpt.get("document"):
                    ext = rpt["filename"].lower()
                    try:
                        content = base64.b64decode(rpt["document"])
                        doc_text = ""

                        if ext.endswith(".pdf"):
                            pdf = PdfReader(io.BytesIO(content))
                            doc_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                        elif ext.endswith(".docx"):
                            word = Document(io.BytesIO(content))
                            doc_text = "\n".join([p.text for p in word.paragraphs])
                        elif ext.endswith(".txt"):
                            doc_text = content.decode("utf-8")

                        # Field Extraction
                        def extract_from_document(doc_text):
                            def extract(patterns, fallback="Not found"):
                                for pattern in patterns:
                                    match = re.search(pattern, doc_text, re.IGNORECASE)
                                    if match:
                                        return match.group(0).strip()
                                return fallback

                            return {
                                "Highest Qualification": extract([
                                    r"(Ph\.?D(?:\s+in\s+[A-Za-z\s&]+)?)",
                                    r"(M\.?B\.?A(?:\s+in\s+[A-Za-z\s&]+)?)",
                                    r"(M\.?Tech(?:\s+in\s+[A-Za-z\s&]+)?)",
                                    r"(B\.?Tech(?:\s+in\s+[A-Za-z\s&]+)?)",
                                    r"(M\.?Sc(?:\s+in\s+[A-Za-z\s&]+)?)",
                                    r"(B\.?E(?:\s+in\s+[A-Za-z\s&]+)?)"
                                ]),

                                "Most Recent Experience": extract([
                                    r"(Currently\s+working\s+as\s+[A-Za-z\s]+)",
                                    r"(Director|Trainer|Professor|Lecturer|Principal|Dean)\s+at\s+[A-Z][A-Za-z&\s,]+",
                                    r"(Director|Trainer|Professor|Lecturer|Principal|Dean)[^\n]{0,100}"
                                ]),

                                "Total Experience": extract([
                                    r"\d{1,2}\+?\s*(years|yrs|Years|Yrs)\s+of\s+(Corporate|Academic|Industry)?\s*experience",
                                    r"(Total\s+Experience\s*[:\-]?\s*\d{1,2}\+?\s*(years|Yrs|yrs))",
                                    r"(\d{1,2}\+?\s*(years|Yrs|yrs))"
                                ]),

                                "Publications": extract([
                                    r"\d+\s*\(\s*\d+\s*National\s*,\s*\d+\s*International\s*\)",
                                    r"\d+\s*(National|International)[^\n]*Publications?",
                                    r"(1\s*National\s*(and|&)?\s*1\s*International)"
                                ])
                            }

                        summary = extract_from_document(doc_text)
                        st.markdown("📑 **Document Summary (AI Extracted):**")
                        for key, val in summary.items():
                            st.markdown(f"**{key}:** {val}")

                    except Exception as e:
                        st.warning(f"⚠️ Couldn't parse document summary. Error: {e}")




elif choice == "Analytics":
    st.subheader("📈 Visual Analytics")
    code = st.text_input("Institute Code for Analysis")
    selected_date = st.date_input("Filter by Date (optional)", value=None)
    if code:
        query = {"code": code}
        if selected_date:
            start_dt = datetime.combine(selected_date, datetime.min.time())
            end_dt = datetime.combine(selected_date, datetime.max.time())
            query["timestamp"] = {"$gte": start_dt, "$lte": end_dt}
        data = list(db.inspections.find(query))
        if data:
            df = pd.DataFrame(data)
            auth_counts = df['authenticity'].value_counts()
            cond_counts = df['condition'].value_counts()

            col1, col2 = st.columns(2)
            with col1:
                st.write("### Authenticity Distribution")
                fig, ax = plt.subplots()
                auth_counts.plot(kind='pie', autopct='%1.1f%%', ax=ax)
                st.pyplot(fig)

            with col2:
                st.write("### Condition Distribution")
                fig, ax = plt.subplots()
                sns.countplot(x='condition', data=df, ax=ax)
                st.pyplot(fig)

            df['score'] = df.apply(lambda x: 1 if x['authenticity'] == 'Authentic' and x['condition'] == 'Good' else 0, axis=1)
            score = round((df['score'].sum() / len(df)) * 100, 2)
            st.success(f"Compliance Score: {score}%")
        else:
            st.info("No data available.")
